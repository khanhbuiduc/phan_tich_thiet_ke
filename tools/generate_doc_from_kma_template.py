from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "example.docx"
OUTPUT = ROOT / "bao_cao_theo_mau_kma.docx"


def clear_document(doc: Document) -> None:
    body = doc.element.body
    sect_pr = body.sectPr
    for child in list(body):
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def find_style(doc: Document, style_id: str, fallback: str) -> str:
    for style in doc.styles:
        if style.style_id == style_id:
            return style.name
    return fallback


def set_page_number(section, start: int, fmt: str | None = None) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))
    if fmt:
        pg_num_type.set(qn("w:fmt"), fmt)
    elif qn("w:fmt") in pg_num_type.attrib:
        del pg_num_type.attrib[qn("w:fmt")]


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(separate)
    run._r.append(text)
    run._r.append(end)


def reset_footer(section) -> None:
    section.footer.is_linked_to_previous = False
    footer = section.footer
    for paragraph in footer.paragraphs:
        p = paragraph._element
        p.getparent().remove(p)
    paragraph = footer.add_paragraph(style="Footer")
    add_page_field(paragraph)


def add_static_toc_lines(doc: Document, items: list[tuple[str, str]]) -> None:
    for style, text in items:
        doc.add_paragraph(text, style=style)


def add_cover(doc: Document, cover_style: str) -> None:
    table = doc.add_table(rows=8, cols=1)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        row.cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    table.cell(0, 0).width = Emu(6076950)

    blocks = [
        [
            "BAN CƠ YẾU CHÍNH PHỦ",
            "HỌC VIỆN KỸ THUẬT MẬT MÃ",
            "¯¯¯¯¯¯¯¯¯¯¯¯¯¯¯¯",
        ],
        [
            "BÁO CÁO PHÂN TÍCH KỸ THUẬT",
            "HỆ THỐNG TRỢ LÝ ẢO TƯ VẤN TUYỂN SINH",
        ],
        [
            "Ngành: Công nghệ thông tin",
            "Mã số: DEMO-2026",
            "Sinh viên thực hiện:",
            "Nhóm phát triển tài liệu",
            "Lớp: Demo-AI",
            "Người hướng dẫn 1:",
            "OpenAI Codex",
            "Khoa/Đơn vị: Mô phỏng theo mẫu KMA",
        ],
        ["Hà Nội, 2026"],
    ]

    all_rows = blocks + deepcopy(blocks)
    all_rows[6][4] = "Lớp: Demo-Word"

    for row_idx, lines in enumerate(all_rows):
        cell = table.cell(row_idx, 0)
        cell.text = ""
        for line_idx, line in enumerate(lines):
            p = cell.paragraphs[0] if line_idx == 0 else cell.add_paragraph()
            p.style = cover_style
            if row_idx in (2, 6):
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.left_indent = Cm(9)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            if row_idx in (1, 5):
                run.bold = True
            if row_idx in (0, 4) and line != "¯¯¯¯¯¯¯¯¯¯¯¯¯¯¯¯":
                run.bold = True
        if row_idx in (3, 7):
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_front_matter(doc: Document, uh1: str, uh2: str, toc1: str, toc2: str, toc3: str) -> None:
    doc.add_page_break()
    doc.add_paragraph("Mục lục", style="Mục lục")
    note = doc.add_paragraph(style="Normal")
    note.add_run(
        "(Tài liệu này được sinh tự động theo bố cục mẫu để kiểm tra khả năng dựng báo cáo Word.)"
    ).bold = True
    add_static_toc_lines(
        doc,
        [
            (toc1, "Một số lời khuyên về trình bày báo cáo kỹ thuật\tii"),
            (toc2, "Về hình thức trình bày tài liệu\tii"),
            (toc2, "Về hình vẽ và bảng biểu\tii"),
            (toc2, "Về phần kết luận\tii"),
            (toc1, "Lời nói đầu\tiii"),
            (toc1, "Chương 1. Khảo sát và phân tích yêu cầu hệ thống\t1"),
            (toc2, "1.1.\tBối cảnh bài toán\t1"),
            (toc2, "1.2.\tMục tiêu hệ thống\t2"),
            (toc1, "Chương 2. Phân tích và thiết kế hệ thống\t4"),
            (toc2, "2.1.\tKiến trúc tổng thể\t4"),
            (toc3, "2.1.1.\tVai trò của chatbot Rasa\t5"),
            (toc3, "2.1.2.\tVai trò của phân hệ RAG\t6"),
            (toc1, "Kết luận\t8"),
            (toc1, "Tài liệu tham khảo\t9"),
        ],
    )

    doc.add_page_break()
    doc.add_paragraph("Lời nói đầu", style=uh1)
    doc.add_paragraph(
        "Tài liệu này được sinh thử để kiểm tra khả năng tái sử dụng mẫu báo cáo học thuật KMA cho đề tài phân tích kỹ thuật hệ thống trợ lý ảo tư vấn tuyển sinh.",
        style="Normal",
    )
    doc.add_paragraph(
        "Mục tiêu của bản sinh thử là giữ đúng phong cách trình bày của mẫu gốc, đồng thời thay thế nội dung bằng một bộ nội dung kỹ thuật mới có cấu trúc rõ ràng và có thể tiếp tục mở rộng về sau.",
        style="Normal",
    )
    doc.add_paragraph("Một số điểm bám theo mẫu", style=uh2)
    doc.add_paragraph(
        "Tài liệu sử dụng cùng hệ lề, cùng phong cách heading, cùng cách đặt tên hình và tên bảng, đồng thời giữ cơ chế đánh số trang theo từng phần giống tài liệu mẫu.",
        style="Normal",
    )


def add_image(doc: Document, image_path: Path, caption_style: str, caption_text: str, width_cm: float = 13.5) -> None:
    p = doc.add_paragraph(style="Center")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    cap = doc.add_paragraph(caption_text, style=caption_style)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_table_caption(doc: Document, style_name: str, caption_text: str) -> None:
    p = doc.add_paragraph(caption_text, style=style_name)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def add_content_table(doc: Document, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, header in enumerate(rows[0]):
        cell = table.rows[0].cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.style = "Table"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
    for row_data in rows[1:]:
        row = table.add_row()
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.style = "Table"
            p.add_run(value)


def add_reference(doc: Document, ref_style: str, text: str) -> None:
    doc.add_paragraph(text, style=ref_style)


def build_document() -> None:
    doc = Document(str(TEMPLATE))
    base_sect_pr = deepcopy(doc.sections[0]._sectPr)
    clear_document(doc)
    doc._body._element.append(base_sect_pr)

    cover_style = find_style(doc, "TrangBa", "Normal")
    uh1 = find_style(doc, "UH1", "Heading 1")
    uh2 = find_style(doc, "UH2", "Heading 2")
    toc1 = find_style(doc, "TOC1", "toc 1")
    toc2 = find_style(doc, "TOC2", "toc 2")
    toc3 = find_style(doc, "TOC3", "toc 3")
    figure_caption = find_style(doc, "Tnhnhv", "Caption")
    table_caption = find_style(doc, "Tnbng", "Caption")
    ref_style = find_style(doc, "Tiliuthamkho", "Normal")

    add_cover(doc, cover_style)

    section2 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_number(section2, 1, "lowerRoman")
    reset_footer(section2)
    add_front_matter(doc, uh1, uh2, toc1, toc2, toc3)

    section3 = doc.add_section(WD_SECTION.NEW_PAGE)
    set_page_number(section3, 1)
    reset_footer(section3)

    doc.add_paragraph("Khảo sát và phân tích yêu cầu hệ thống", style="Heading 1")
    doc.add_paragraph("Bối cảnh bài toán", style="Heading 2")
    doc.add_paragraph(
        "Trong giai đoạn tuyển sinh, thí sinh và phụ huynh cần tra cứu nhanh những thông tin như ngành đào tạo, phương thức xét tuyển, học phí, điểm chuẩn và quy trình nhập học. Khi thông tin nằm rải rác trên nhiều kênh khác nhau, bộ phận tư vấn phải tiếp nhận nhiều câu hỏi lặp lại và khó duy trì sự nhất quán trong phản hồi.",
        style="Normal",
    )
    doc.add_paragraph(
        "Vì vậy, hệ thống trợ lý ảo được đề xuất nhằm tự động hóa các câu hỏi thường gặp, hỗ trợ hội thoại tự nhiên và truy xuất dữ liệu chính thống từ tập tài liệu tuyển sinh của nhà trường.",
        style="Normal",
    )
    add_image(
        doc,
        ROOT / "content" / "source" / "Hình 2.1 Mô hình tư vấn tuyển sinh truyền thống và mô hình có chatbot hỗ trợ..png",
        figure_caption,
        "Hình 1.1. Mô hình tư vấn tuyển sinh truyền thống và mô hình có chatbot hỗ trợ",
    )

    doc.add_paragraph("Mục tiêu hệ thống", style="Heading 2")
    doc.add_paragraph(
        "Hệ thống được thiết kế để phản hồi nhanh, trả lời đúng theo nguồn dữ liệu đã kiểm duyệt và cho phép cập nhật tri thức tuyển sinh theo từng mùa. Ngoài ra, phân hệ quản trị phải hỗ trợ theo dõi câu hỏi khó, cải tiến dữ liệu và đánh giá chất lượng phản hồi.",
        style="Normal",
    )
    add_table_caption(doc, table_caption, "Bảng 1.1. Mục tiêu chính của hệ thống trợ lý ảo tư vấn tuyển sinh")
    add_content_table(
        doc,
        [
            ["Nhóm mục tiêu", "Mô tả"],
            ["Hiệu quả tư vấn", "Giảm tải cho bộ phận tư vấn và rút ngắn thời gian phản hồi."],
            ["Độ chính xác", "Bám theo thông tin tuyển sinh chính thống và nguồn dữ liệu đã kiểm duyệt."],
            ["Mở rộng", "Dễ cập nhật ngành học, biểu mẫu, phương thức xét tuyển và tài liệu mới."],
            ["Theo dõi chất lượng", "Lưu vết hội thoại và ghi nhận những câu hỏi cần cải thiện thêm."],
        ],
    )

    doc.add_paragraph("Phân tích và thiết kế hệ thống", style="Heading 1")
    doc.add_paragraph("Kiến trúc tổng thể", style="Heading 2")
    doc.add_paragraph(
        "Kiến trúc của hệ thống được chia thành giao diện người dùng, backend điều phối, chatbot Rasa, action server, phân hệ RAG và các lớp lưu trữ dữ liệu. Cách chia lớp này cho phép xử lý tốt cả những câu hỏi nghiệp vụ lặp lại lẫn những câu hỏi cần bám theo tài liệu dài.",
        style="Normal",
    )
    add_image(
        doc,
        ROOT / "content" / "source" / "Hình 3.1 Sơ đồ kiến trúc tổng thể hệ thống..png",
        figure_caption,
        "Hình 2.1. Sơ đồ kiến trúc tổng thể hệ thống",
    )

    doc.add_paragraph("Vai trò của chatbot Rasa", style="Heading 3")
    doc.add_paragraph(
        "Rasa chịu trách nhiệm phân tích ý định, nhận diện thực thể và điều phối các luồng hội thoại có cấu trúc. Phân hệ này phù hợp với các câu hỏi như học phí, phương thức xét tuyển, điểm chuẩn và thủ tục nhập học.",
        style="Normal",
    )
    doc.add_paragraph("Vai trò của phân hệ RAG", style="Heading 3")
    doc.add_paragraph(
        "RAG hỗ trợ truy xuất tri thức từ kho tài liệu tuyển sinh, biểu mẫu và văn bản nghiệp vụ. Khi câu hỏi vượt ra ngoài các intent cố định, hệ thống có thể sử dụng truy xuất ngữ nghĩa để tổng hợp câu trả lời bám sát tài liệu gốc.",
        style="Normal",
    )
    add_image(
        doc,
        ROOT / "content" / "source" / "Hình 3.6 Query Flow của RAG..png",
        figure_caption,
        "Hình 2.2. Query Flow của phân hệ RAG",
    )

    doc.add_paragraph("So sánh các hướng xử lý", style="Heading 2")
    add_table_caption(doc, table_caption, "Bảng 2.1. So sánh các hướng xử lý chính trong hệ thống")
    add_content_table(
        doc,
        [
            ["Thành phần", "Điểm mạnh", "Phạm vi áp dụng"],
            ["Rule-based", "Đơn giản, dễ kiểm soát", "FAQ cố định hoặc luồng xác định rõ"],
            ["Rasa", "Hiểu intent và quản lý hội thoại tốt", "Nghiệp vụ tuyển sinh có cấu trúc"],
            ["RAG", "Truy xuất tri thức từ tài liệu dài", "Câu hỏi cần bám theo văn bản chính thống"],
        ],
    )

    doc.add_paragraph("Kết luận", style=uh1)
    doc.add_paragraph(
        "Bản sinh thử này cho thấy có thể tạo một tài liệu Word mới theo đúng phong cách trình bày của mẫu KMA, đồng thời thay thế nội dung gốc bằng nội dung báo cáo kỹ thuật cho hệ thống trợ lý ảo tư vấn tuyển sinh. Cấu trúc hiện tại đủ làm nền để tiếp tục mở rộng thành bản chuyển đổi đầy đủ từ bộ nguồn LaTeX hiện có.",
        style="Normal",
    )

    doc.add_paragraph("Tài liệu tham khảo", style=uh1)
    add_reference(doc, ref_style, "[1]. Tài liệu phân tích kỹ thuật nội bộ của dự án trợ lý ảo tư vấn tuyển sinh.")
    add_reference(doc, ref_style, "[2]. Rasa Documentation, nền tảng xây dựng chatbot hội thoại có kiểm soát.")
    add_reference(doc, ref_style, "[3]. Tài liệu kiến trúc và luồng truy xuất tri thức của phân hệ RAG trong thư mục content/source.")

    doc.core_properties.title = "Báo cáo theo mẫu KMA"
    doc.core_properties.subject = "Sinh doc theo mẫu KMA"
    doc.core_properties.author = "OpenAI Codex"
    doc.save(str(OUTPUT))


if __name__ == "__main__":
    build_document()
